import docx
import csv
import sys
import os

def extract_glossary_table(input_path, output_path):
    print(f"Reading {input_path}...")
    try:
        doc = docx.Document(input_path)
        
        if not doc.tables:
            print("No tables found in the document.")
            return

        # Assume the first table is the glossary table
        table = doc.tables[0]
        data = []
        
        # Determine header row and start index for data
        header_row_text = [cell.text.strip() for cell in table.rows[0].cells]
        
        en_col_index = -1
        zh_col_index = -1
        
        # Find the English and Chinese columns based on headers
        for i, header in enumerate(header_row_text):
            if "English" in header or "原文术语" in header:
                en_col_index = i
            if "Chinese" in header or "推荐译文" in header:
                zh_col_index = i
        
        if en_col_index == -1 or zh_col_index == -1:
            print("Could not find 'English' and 'Chinese' columns in the table header.")
            return

        print("Extracting terms from the first table...")
        
        # Iterate through rows, skipping the header
        for i, row in enumerate(table.rows):
            if i == 0: # Skip header row
                continue

            cells = row.cells
            
            # Ensure row has enough columns
            if len(cells) > max(en_col_index, zh_col_index):
                en_term = cells[en_col_index].text.strip()
                zh_term = cells[zh_col_index].text.strip()
                
                if en_term and zh_term: # Only add if both terms are present
                    data.append({
                        "EN_TERM": en_term,
                        "ZH_TERM": zh_term
                    })
            
        print(f"Extracted {len(data)} terms.")
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=["EN_TERM", "ZH_TERM"], delimiter='\t')
            writer.writeheader()
            writer.writerows(data)
            
        print(f"Written to {output_path}")

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_glossary_table.py <input_docx> <output_tsv>")
    else:
        extract_glossary_table(sys.argv[1], sys.argv[2])
