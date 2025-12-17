import docx
import sys
import os

def extract_table_data(doc_path):
    """Extracts data from the first table of a docx document."""
    try:
        doc = docx.Document(doc_path)
        if not doc.tables:
            return None, "No tables found."
        
        table = doc.tables[0]
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        return data, None
    except Exception as e:
        return None, f"Error processing {doc_path}: {e}"

def compare_docx_files(file1_path, file2_path):
    print(f"Comparing '{os.path.basename(file1_path)}' and '{os.path.basename(file2_path)}'\n")

    data1, error1 = extract_table_data(file1_path)
    data2, error2 = extract_table_data(file2_path)

    if error1:
        print(f"Error reading file 1: {error1}")
        return
    if error2:
        print(f"Error reading file 2: {error2}")
        return
    
    if data1 is None and data2 is None:
        print("Both documents have no tables to compare.")
        return
    if data1 is None:
        print("File 1 has no tables, File 2 has tables.")
        return
    if data2 is None:
        print("File 2 has no tables, File 1 has tables.")
        return

    print("--- Table Comparison (First Table Only) ---")
    print(f"File 1 (Old): Rows = {len(data1)}, Cols = {len(data1[0]) if data1 else 0}")
    print(f"File 2 (New): Rows = {len(data2)}, Cols = {len(data2[0]) if data2 else 0}")

    if len(data1) != len(data2):
        print(f"DIFFERENCE: Number of rows are different (File 1: {len(data1)}, File 2: {len(data2)})")
    if data1 and data2 and len(data1[0]) != len(data2[0]):
        print(f"DIFFERENCE: Number of columns are different (File 1: {len(data1[0])}, File 2: {len(data2[0])})")

    # Compare cell by cell
    differences_found = False
    max_diffs_to_show = 5
    diff_count = 0

    min_rows = min(len(data1), len(data2))
    min_cols = min(len(data1[0]), len(data2[0])) if data1 and data2 else 0

    for r in range(min_rows):
        for c in range(min_cols):
            cell1 = data1[r][c]
            cell2 = data2[r][c]
            if cell1 != cell2:
                differences_found = True
                diff_count += 1
                if diff_count <= max_diffs_to_show:
                    print(f"\nDIFFERENCE in Row {r}, Col {c}:")
                    print(f"  File 1: '{cell1[:100]}...'") # Show first 100 chars
                    print(f"  File 2: '{cell2[:100]}...'")
                elif diff_count == max_diffs_to_show + 1:
                    print("\n(Further differences suppressed...)")
                break # Move to next row after first difference in a row

    if not differences_found:
        print("\nSIMILARITY: No cell-by-cell differences found in common rows/columns.")
        if len(data1) == len(data2) and (not data1 or len(data1[0]) == len(data2[0])):
            print("SIMILARITY: Both tables (first table) are identical in content and structure.")
        else:
            print("SIMILARITY: Common parts are identical, but structure (rows/cols) differs.")
    else:
        print(f"\nTotal significant cell differences detected: {diff_count}")
        print("DIFFERENCE: Content differs between the two files (first table).")

    print("\n--- Summary ---")
    if differences_found or len(data1) != len(data2) or (data1 and data2 and len(data1[0]) != len(data2[0])):
        print("DIFFERENCE: The two .docx files are NOT identical in their first table's content or structure.")
    else:
        print("SIMILARITY: The two .docx files appear to be identical in their first table's content and structure.")


if __name__ == "__main__":
    file1 = "sample/Framework Contract for Supply of Modules-TW250716 适用于通威股份-(离线_纯文本_25121600424).docx"
    file2 = "sample/Framework Contract for Supply of Modules-TW250716 适用于通威股份-(离线_纯文本无锁重_25121600424).docx"
    compare_docx_files(file1, file2)
