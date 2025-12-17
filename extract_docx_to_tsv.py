import docx
import csv
import sys
import os

def extract_docx_to_tsv(input_path, output_path):
    print(f"Reading {input_path}...")
    try:
        doc = docx.Document(input_path)
        
        if not doc.tables:
            print("No tables found in the document.")
            return

        table = doc.tables[0]
        data = []
        
        # Headers: ID, EN, ZH
        # Assuming the doc structure: ID(0, messy), #(1, clean ID), Source(2, EN), Match(3), Target(4, ZH)
        
        print("Extracting rows...")
        rows = table.rows
        # Skip header if it matches known headers
        start_index = 0
        if len(rows) > 0:
            header_cells = [c.text.strip() for c in rows[0].cells]
            if "Source" in header_cells and "Target" in header_cells:
                start_index = 1
        
        for row in rows[start_index:]:
            cells = row.cells
            if len(cells) < 5:
                continue
                
            row_id = cells[1].text.strip()
            en_text = cells[2].text.strip()
            zh_text = cells[4].text.strip()
            
            # Simple validation to avoid empty rows if any
            if not row_id and not en_text and not zh_text:
                continue

            data.append({
                "ID": row_id,
                "EN": en_text,
                "ZH": zh_text
            })
            
        print(f"Extracted {len(data)} rows.")
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=["ID", "EN", "ZH"], delimiter='\t')
            writer.writeheader()
            writer.writerows(data)
            
        print(f"Written to {output_path}")

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_docx_to_tsv.py <input_docx> <output_tsv>")
    else:
        extract_docx_to_tsv(sys.argv[1], sys.argv[2])
