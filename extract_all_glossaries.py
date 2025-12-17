import docx
import csv
import sys
import re

def normalize_text(text):
    if not text:
        return ""
    return text.strip()

def find_column_index(headers, keywords):
    for idx, header in enumerate(headers):
        for kw in keywords:
            if kw.lower() in header.lower():
                return idx
    return -1

def extract_all_glossaries(input_path, output_path):
    print(f"Reading {input_path}...")
    try:
        doc = docx.Document(input_path)
        all_terms = []
        
        # Keywords for identifying columns
        en_keywords = ['english', '原文', 'source', 'class', '英文缩写', 'share'] 
        # Added 'share' just in case, but 'class' covers '股份类别 (Class)'
        # '英文缩写' covers Table 5.
        
        zh_keywords = ['chinese', '译文', 'target', 'translation', '推荐译名', '推荐中文', '中文']

        print(f"Found {len(doc.tables)} tables.")

        for i, table in enumerate(doc.tables):
            rows = table.rows
            if not rows:
                continue
                
            # Assume first row is header
            headers = [cell.text.strip() for cell in rows[0].cells]
            
            en_idx = find_column_index(headers, en_keywords)
            zh_idx = find_column_index(headers, zh_keywords)
            
            # Fallback for Table 5 if "英文缩写" isn't caught by general keys? 
            # I added '英文缩写' to en_keywords, so it should be caught.
            
            if en_idx == -1 or zh_idx == -1:
                print(f"Table {i+1}: Could not identify EN/ZH columns. Headers: {headers}")
                # Fallback logic based on table structure if headers fail?
                # For now, skip.
                continue
            
            print(f"Table {i+1}: Extracting EN (Col {en_idx}) -> ZH (Col {zh_idx}) | Headers: {headers[en_idx]} -> {headers[zh_idx]}")
            
            for j, row in enumerate(rows[1:]): # Skip header
                cells = row.cells
                if len(cells) <= max(en_idx, zh_idx):
                    continue
                
                en_term = normalize_text(cells[en_idx].text)
                zh_term = normalize_text(cells[zh_idx].text)
                
                # Clean up: remove newlines
                en_term = en_term.replace('\n', ' ')
                zh_term = zh_term.replace('\n', ' ')
                
                if en_term and zh_term:
                    all_terms.append((en_term, zh_term))

        # Deduplicate while preserving order
        unique_terms = []
        seen = set()
        for en, zh in all_terms:
            # Create a unique key. Maybe strictly EN is enough? 
            # Or (EN, ZH). Let's use EN to avoid duplicates with different translations unless intended.
            # Usually glossary keys are unique.
            if en.lower() not in seen:
                seen.add(en.lower())
                unique_terms.append({'EN_TERM': en, 'ZH_TERM': zh})
        
        print(f"Extracted {len(unique_terms)} unique terms from {len(all_terms)} total rows.")
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=["EN_TERM", "ZH_TERM"], delimiter='\t')
            writer.writeheader()
            writer.writerows(unique_terms)
            
        print(f"Written to {output_path}")

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_all_glossaries.py <input_docx> <output_tsv>")
    else:
        extract_all_glossaries(sys.argv[1], sys.argv[2])
